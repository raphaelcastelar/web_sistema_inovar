from __future__ import annotations

import re
from datetime import datetime
from decimal import Decimal, InvalidOperation, ROUND_HALF_UP
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


INSS_ALIQUOTA = Decimal("0.11")
INSS_TETO_BASE = Decimal("8475.55")
IR_FAIXA_ISENCAO_TOTAL = Decimal("5000.00")
IR_FAIXA_REDUCAO = Decimal("7350.00")
IR_REDUCAO_INTERCEPT = Decimal("978.62")
IR_REDUCAO_SLOPE = Decimal("0.133145")


def _round_currency(value: Decimal) -> Decimal:
    return value.quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)


def _to_decimal(value, field_name: str, required: bool = True) -> Decimal:
    if value in (None, ""):
        if required:
            raise ValueError(f"O campo '{field_name}' e obrigatorio.")
        return Decimal("0.00")

    normalized = str(value).strip()
    if "," in normalized:
        normalized = normalized.replace(".", "").replace(",", ".")
    else:
        normalized = normalized.replace(",", "")
    try:
        return Decimal(normalized)
    except (InvalidOperation, ValueError):
        raise ValueError(f"O campo '{field_name}' deve ser numerico.")


def _calc_irrf_sem_reducao(base_calculo: Decimal) -> Decimal:
    if base_calculo <= Decimal("2428.80"):
        return Decimal("0.00")
    if base_calculo <= Decimal("2826.65"):
        return _round_currency((base_calculo * Decimal("0.075")) - Decimal("182.16"))
    if base_calculo <= Decimal("3751.05"):
        return _round_currency((base_calculo * Decimal("0.15")) - Decimal("394.16"))
    if base_calculo <= Decimal("4664.68"):
        return _round_currency((base_calculo * Decimal("0.225")) - Decimal("675.49"))
    return _round_currency((base_calculo * Decimal("0.275")) - Decimal("908.73"))


def _calcular_impostos_pro_labore(valor_bruto: Decimal) -> tuple[Decimal, Decimal, Decimal]:
    bruto = max(valor_bruto, Decimal("0.00"))

    base_inss = min(bruto, INSS_TETO_BASE)
    valor_inss = _round_currency(base_inss * INSS_ALIQUOTA)

    base_irrf = max(Decimal("0.00"), bruto - valor_inss)
    irrf_sem_reducao = max(Decimal("0.00"), _calc_irrf_sem_reducao(base_irrf))

    if base_irrf <= IR_FAIXA_ISENCAO_TOTAL:
        valor_irrf = Decimal("0.00")
    elif base_irrf <= IR_FAIXA_REDUCAO:
        reducao = max(
            Decimal("0.00"),
            IR_REDUCAO_INTERCEPT - (IR_REDUCAO_SLOPE * base_irrf),
        )
        valor_irrf = max(Decimal("0.00"), irrf_sem_reducao - reducao)
    else:
        valor_irrf = irrf_sem_reducao

    valor_irrf = _round_currency(valor_irrf)
    valor_liquido = _round_currency(max(Decimal("0.00"), bruto - valor_inss - valor_irrf))
    return valor_inss, valor_irrf, valor_liquido


def _format_currency_br(value: Decimal) -> str:
    value = _round_currency(value)
    inteiro, decimal = f"{value:.2f}".split(".")
    inteiro = f"{int(inteiro):,}".replace(",", ".")
    return f"R$ {inteiro},{decimal}"


def _required_text(data: dict, key: str, label: str) -> str:
    value = str(data.get(key, "")).strip()
    if not value:
        raise ValueError(f"O campo '{label}' e obrigatorio.")
    return value


def _optional_text(data: dict, key: str, default: str = "") -> str:
    return str(data.get(key, default)).strip() or default


def _default_filename(nome_colaborador: str, referencia: str) -> str:
    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", nome_colaborador.strip()).strip("_").lower()
    safe_ref = re.sub(r"[^0-9]+", "_", referencia).strip("_")
    if not safe_name:
        safe_name = "colaborador"
    if not safe_ref:
        safe_ref = datetime.now().strftime("%m_%Y")
    return f"recibo_pro_labore_{safe_name}_{safe_ref}.docx"


def build_pro_labore_docx(payload: dict) -> tuple[bytes, str]:
    empresa_nome = _required_text(payload, "empresa_nome", "empresa_nome")
    empresa_endereco = _required_text(payload, "empresa_endereco", "empresa_endereco")
    empresa_numero = _required_text(payload, "empresa_numero", "empresa_numero")
    empresa_bairro = _required_text(payload, "empresa_bairro", "empresa_bairro")
    empresa_municipio = _required_text(payload, "empresa_municipio", "empresa_municipio")
    empresa_estado = _required_text(payload, "empresa_estado", "empresa_estado")
    empresa_cep = _required_text(payload, "empresa_cep", "empresa_cep")
    empresa_cnpj = _required_text(payload, "empresa_cnpj", "empresa_cnpj")

    colaborador_nome = _required_text(payload, "colaborador_nome", "colaborador_nome")
    colaborador_cpf = _required_text(payload, "colaborador_cpf", "colaborador_cpf")
    referencia_mes_ano = _required_text(payload, "referencia_mes_ano", "referencia_mes_ano")
    data_assinatura = _required_text(payload, "data_assinatura", "data_assinatura")
    local_assinatura = _required_text(payload, "local_assinatura", "local_assinatura")
    valor_liquido_extenso = _required_text(payload, "valor_liquido_extenso", "valor_liquido_extenso")

    valor_bruto = _to_decimal(payload.get("valor_bruto"), "valor_bruto")
    valor_inss_calculado, valor_irrf_calculado, valor_liquido_calculado = _calcular_impostos_pro_labore(valor_bruto)
    valor_inss_payload = payload.get("valor_inss")
    valor_irrf_payload = payload.get("valor_irrf")
    valor_liquido_payload = payload.get("valor_liquido")
    valor_inss = (
        _to_decimal(valor_inss_payload, "valor_inss", required=False)
        if valor_inss_payload not in (None, "")
        else valor_inss_calculado
    )
    valor_irrf = (
        _to_decimal(valor_irrf_payload, "valor_irrf", required=False)
        if valor_irrf_payload not in (None, "")
        else valor_irrf_calculado
    )
    valor_liquido = (
        _to_decimal(valor_liquido_payload, "valor_liquido", required=False)
        if valor_liquido_payload not in (None, "")
        else valor_liquido_calculado
    )

    doc = Document()

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("RECIBO DE PRO-LABORE")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)

    doc.add_paragraph()

    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "Verba"
    cabecalho[1].text = "Referencia"
    cabecalho[2].text = "Vencimentos"
    cabecalho[3].text = "Descontos"

    linha_bruto = tabela.add_row().cells
    linha_bruto[0].text = "0702 - Retirada Pro-Labore Diretor"
    linha_bruto[1].text = referencia_mes_ano
    linha_bruto[2].text = _format_currency_br(valor_bruto)
    linha_bruto[3].text = "-"

    linha_inss = tabela.add_row().cells
    linha_inss[0].text = "0526 - INSS Contribuinte Individual"
    linha_inss[1].text = referencia_mes_ano
    linha_inss[2].text = "-"
    linha_inss[3].text = _format_currency_br(valor_inss)

    linha_irrf = tabela.add_row().cells
    linha_irrf[0].text = "0530 - Desconto de IRRF"
    linha_irrf[1].text = referencia_mes_ano
    linha_irrf[2].text = "-"
    linha_irrf[3].text = _format_currency_br(valor_irrf)

    doc.add_paragraph()
    p_liquido = doc.add_paragraph()
    p_liquido.add_run("Liquido Recebido: ").bold = True
    p_liquido.add_run(_format_currency_br(valor_liquido))

    texto = (
        f"Recebi de {empresa_nome}, sediada no {empresa_endereco}, n. {empresa_numero}, "
        f"bairro {empresa_bairro}, municipio de {empresa_municipio}, estado {empresa_estado}, "
        f"CEP {empresa_cep}, CNPJ {empresa_cnpj}, a importancia supra de "
        f"{_format_currency_br(valor_liquido)} ({valor_liquido_extenso}) referente ao meu "
        f"Pro-Labore do mes {referencia_mes_ano}, com os descontos exigidos em lei. "
        "Para maior clareza e devidos fins de direito, firmo o presente."
    )
    doc.add_paragraph(texto)

    doc.add_paragraph()
    p_local_data = doc.add_paragraph(f"{local_assinatura}, {data_assinatura}.")
    p_local_data.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    assinatura = doc.add_paragraph("________________________________________")
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome = doc.add_paragraph(colaborador_nome)
    nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cpf = doc.add_paragraph(f"CPF: {colaborador_cpf}")
    cpf.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = _optional_text(payload, "nome_arquivo")
    if not filename:
        filename = _default_filename(colaborador_nome, referencia_mes_ano)
    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    return output.getvalue(), filename
